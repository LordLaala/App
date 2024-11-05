import json
import os
from pathlib import Path
import docx
from future.backports.datetime import datetime

import OCR_Services.OCR as OCR
from DocConverter import DocConverter

from dateutil.parser import parse
import datefinder


class JsonConverter:
    """
    Initialize the JsonConverter with the given file path and image directory.

    :param filepath: Path to the .docx file.
    :param image_dir: Directory to save extracted images.
    """
    def __init__(self, filepath, image_dir="images"):
        self.path = Path(filepath)
        self.doc = docx.Document(self.path)
        self.image_names = []
        self.image_dir = image_dir
        os.makedirs(self.image_dir, exist_ok=True)
        self.data = self.create_data_from_docx(self.doc)

    @staticmethod
    def is_date(string):
        """
        Check if the given string is a date.

        :param string: String to check.
        :return: True if the string is a date, False otherwise.
        """
        try:
            parse(string, fuzzy=True, dayfirst=True)
            return True
        except ValueError:
            return False

    def get_text_after_tables(self, doc):
        """
        Get the text after the tables in the document.

        :param doc: The docx document object.
        :return: List of text strings found after the tables.
        """
        text_after_tables = []
        table_found = False

        for element in doc.element.body:
            if element.tag.endswith("tbl"):
                table_found = True
            elif table_found and element.tag.endswith("p"):
                text = element.text.strip()
                if text:
                    text_after_tables.append(text)
            elif table_found and not element.tag.endswith("p"):
                table_found = False

        return text_after_tables

    def save_data_to_json(self, file_path):
        """Save the data to a json file.
        :param file_path: The path to the json file.
        :return None
        """
        text = self.get_text_after_tables(self.doc)

        # Alternative to get the text from the document
        content_text = ""
        for content in text:
            if content_text == "":
                content_text = content
                continue
            content_text = f"{content_text}\n{content}"
        self.data["Zusaetzliche Informationen"] = [content_text]

        with open(file_path, "w",encoding="utf-8") as file:
            json.dump(self.data, file, indent=4, ensure_ascii=False)

    def process_table(self, doc, table_index, key, keys, values):
        """
        Process a table in the document and extract keys and values.

        :param doc: The docx document object.
        :param table_index: Index of the table to process.
        :param key: Boolean indicating if the current row contains keys.
        :param keys: List to store extracted keys.
        :param values: List to store extracted values.
        :return: None
        """
        if table_index >= len(doc.tables):
            return

        table = doc.tables[table_index]
        for row in table.columns:
            elements = []
            for cell in row.cells:
                count = 0
                cell_text = cell.text.strip()
                elements.append(cell_text)
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        if run.element.xpath(".//w:drawing"):
                            for drawing in run.element.xpath(".//w:drawing"):
                                image = drawing.xpath(".//a:blip/@r:embed")
                                if image:
                                    for image_id in image:
                                        image_part = doc.part.related_parts[image_id]
                                        image_filename = os.path.join(
                                            self.image_dir, f"{image_id}.png"
                                        )
                                        self.image_names.append(f"{image_id}.png")
                                        with open(image_filename, "wb") as img_file:
                                            count += 1
                                            img_file.write(image_part.blob)
                                            image_text = "Hie ist ein Bild"
                                            """"OCR.get_text_from_image(
                                                f"/home/muster/CoSo/PythonProjects/Word/images/{image_id}.png"
                                            )"""
                                            cell_index = len(elements) - 1
                                            elements[
                                                cell_index
                                            ] = f"{elements[cell_index]} \n<<Screenshot[{count}]:\n{image_text}>> "

            if key:
                keys.extend(elements)
                key = False
            else:
                values.extend(elements)
                key = True

        self.process_table(doc, table_index + 1, key, keys, values)

    def create_data_from_docx(self, doc):
        """
        Create a dictionary of data extracted from the docx document.

        :param doc: The docx document object.
        :return: Dictionary containing the extracted data.
        """
        keys = []
        values = []
        self.process_table(doc, 0, True, keys, values)
        result_data = {}

        for i in range(len(keys) - 1, -1, -1):
            if keys[i] in ["Datum", "Datum:"]:
                break
            if keys[i] != "" and not self.is_date(keys[i]) and i <= len(values):
                try:
                    text = values[i]
                    date = datefinder.find_dates(text=text, first="day", strict=True)
                    dates = []

                    for date in date:
                        dates.append(datetime.strftime(date, "%d.%m.%Y"))
                    if len(dates) == 1:
                        values[i] = f"<{keys[i]}> {values[i]}"
                        keys[i] = dates[0]
                except Exception as e:
                    continue
        do_it = False
        if len(keys) == len(values):
            for i in range(len(values)):
                if values[i] in ['Gemeinde', 'Kat.-Nr. (aktuell)', 'Zuteilung', 'Kat.-Nr. (alt)', 'Kat.-Nr. (alt)', 'Datum',
                                 'Datum:', 'Alt Kat. Nr.']:
                    do_it = True
            if do_it:
                for i in range(len(values)):
                    if i % 2 == 0:
                        continue
                    else:
                        if JsonConverter.is_date(keys[i]):
                            continue
                        temp = values[i]
                        values[i] = keys[i]
                        keys[i] = temp

            for i in range(len(keys)):
                if keys[i] in ['Gemeinde', 'Kat.-Nr. (aktuell)', 'Zuteilung', 'Kat.-Nr. (alt)', 'Kat.-Nr. (alt)', 'Datum',
                             'Datum:', 'Alt Kat. Nr.'] or JsonConverter.is_date(keys[i]):
                    if keys[i] != values[i]:
                        if keys[i] in result_data:
                            result_data[keys[i]] = [f"{result_data[keys[i]][0]} {values[i]}"]
                        else:
                            result_data[keys[i]] = [values[i]]
        else:
            pass
        return result_data
        """    
        else:

            if len(keys)>len(values):
                for i in range(len(keys)):
                    if keys[i] in ['Gemeinde', 'Kat.-Nr. (aktuell)', 'Zuteilung', 'Kat.-Nr. (alt)', 'Kat.-Nr. (alt)',
                                   'Datum',
                                   'Datum:', 'Alt Kat. Nr.'] or JsonConverter.is_date(keys[i]):
                        if i % 2 == 0:
                            if JsonConverter.is_date(keys[i]):
                                if keys[i] not in result_data:
                                    result_data[keys[i]] = []


                                if keys[i] not in result_data:
                                    result_data[keys[i]] = []
                                if i <= len(values):
                                    result_data[keys[i]].append(values[i])
                    else:
                        if i < len(values):
                            if values[i] in ['Gemeinde', 'Kat.-Nr. (aktuell)', 'Zuteilung', 'Kat.-Nr. (alt)', 'Kat.-Nr. (alt)', 'Datum',
                                 'Datum:', 'Alt Kat. Nr.'] or JsonConverter.is_date(values[i]):
                                if i % 2 != 0:
                                    if values[i] not in result_data:
                                        result_data[values[i]] = []
                                    result_data[values[i]].append(keys[i])
                            else:
                                continue
"""


